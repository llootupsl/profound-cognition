#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# 作者：阿洋
"""
build_docx.py —— 从 Markdown 报告构建 .docx 文件（纯 python-docx 实现）
使用场景：当 Pandoc 不可用时的穷尽尝试路径。
"""

import argparse
import datetime
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE


# ============================================================================
# 常量
# ============================================================================

FONT_BODY = "未来荧黑"
FONT_MONO = "Consolas"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_CODE = Pt(9)
FONT_SIZE_COVER_TITLE = Pt(28)
FONT_SIZE_COVER_SUBTITLE = Pt(14)
FONT_SIZE_COVER_META = Pt(11)
LINE_SPACING_BODY = 1.5
LINE_SPACING_CODE = 1.15

COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY_BG = RGBColor(0xF2, 0xF2, 0xF2)
COLOR_QUOTE_BG = RGBColor(0xF5, 0xF5, 0xF5)
COLOR_QUOTE_BORDER = RGBColor(0x66, 0x66, 0x66)
COLOR_CODE_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_TABLE_BORDER = RGBColor(0x00, 0x00, 0x00)
COLOR_CAPTION = RGBColor(0x55, 0x55, 0x55)

# 三线表 border weight (eighths of a point; 1.5pt = 12, 0.75pt = 6)
TABLE_BORDER_TOP_BOTTOM = "12"
TABLE_BORDER_HEADER = "6"

COVER_HEIGHT = Cm(29.7)  # A4 height
COVER_WIDTH = Cm(21.0)   # A4 width


# ============================================================================
# 辅助函数
# ============================================================================

def _set_cell_border(cell, **kwargs):
    """设置单元格边框。kwargs: top, bottom, left, right, insideH, insideV"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:val="single" w:sz="{val["sz"]}" w:space="0" '
            f'w:color="{val.get("color","000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _set_run_font(run, name=FONT_BODY, size=FONT_SIZE_BODY, bold=False,
                  italic=False, color=COLOR_BLACK):
    """设置 run 的字号/字体等"""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def _set_paragraph_spacing(paragraph, line_spacing=LINE_SPACING_BODY,
                           space_before=0, space_after=0):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def _add_shading(paragraph, color_hex="F2F2F2"):
    """给段落加背景色"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" '
        f'w:val="clear"/>'
    )
    pPr.append(shd)


def _add_left_border(paragraph, color_hex="666666", sz="12", space=4):
    """给段落加左侧边框"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="{sz}" w:space="{space}" '
        f'w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ============================================================================
# 覆盖页
# ============================================================================

def add_cover(doc, meta):
    """
    添加封面页。
    meta: dict with keys title, subtitle, author, version, run_date
    """
    # --- 封面节 ---
    section = doc.sections[0]
    section.page_height = COVER_HEIGHT
    section.page_width = COVER_WIDTH
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 顶部留白
    for _ in range(5):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, line_spacing=1.0, space_before=0, space_after=0)

    # 标题
    title = meta.get("title", "未命名报告")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    _set_run_font(run, size=FONT_SIZE_COVER_TITLE, bold=True)
    _set_paragraph_spacing(p_title, line_spacing=1.2, space_before=0, space_after=Pt(12))

    # 副标题
    subtitle = meta.get("subtitle", "")
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run(subtitle)
        _set_run_font(run, size=FONT_SIZE_COVER_SUBTITLE, color=RGBColor(0x44, 0x44, 0x44))
        _set_paragraph_spacing(p_sub, line_spacing=1.2, space_before=0, space_after=Pt(20))

    # 分隔线
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("─" * 40)
    _set_run_font(run, size=Pt(10), color=RGBColor(0xAA, 0xAA, 0xAA))
    _set_paragraph_spacing(p_line, line_spacing=1.0, space_before=0, space_after=Pt(20))

    # 元信息
    meta_lines = []
    run_date = meta.get("run_date", datetime.date.today().strftime("%Y-%m-%d"))
    meta_lines.append(f"日期：{run_date}")

    author = meta.get("author", "")
    if author:
        meta_lines.append(f"作者：{author}")

    version = meta.get("version", "v1.0")
    meta_lines.append(f"版本：{version}")

    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        _set_run_font(run, size=FONT_SIZE_COVER_META, color=RGBColor(0x66, 0x66, 0x66))
        _set_paragraph_spacing(p, line_spacing=1.5, space_before=0, space_after=0)

    # 封面后分页
    doc.add_page_break()


# ============================================================================
# 目录域
# ============================================================================

def add_toc_field(doc):
    """
    插入「目录」字段 —— Word TOC 域代码。
    用户打开文档后按 F9 刷新即可生成目录。
    """
    p_heading = doc.add_paragraph()
    p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_heading.add_run("目  录")
    _set_run_font(run, size=Pt(18), bold=True)
    _set_paragraph_spacing(p_heading, line_spacing=1.5, space_before=0, space_after=Pt(12))

    p_toc = doc.add_paragraph()
    # 插入 TOC 域
    run = p_toc.add_run()
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar_begin)

    run2 = p_toc.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        f' TOC \\o "1-3" \\h \\z \\u '
        f'</w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p_toc.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar_end)

    run4 = p_toc.add_run("[ 请在 Word 中右键此处 → 更新域，以生成目录 ]")
    _set_run_font(run4, size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))

    run5 = p_toc.add_run()
    fldChar_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar_end2)

    _set_paragraph_spacing(p_toc, line_spacing=1.0, space_before=0, space_after=Pt(6))

    doc.add_page_break()


# ============================================================================
# 块渲染
# ============================================================================

def render_block(doc, block):
    """
    渲染一个内容块。
    block: dict with 'type' key:
        - heading: {type, level, text}
        - paragraph: {type, text, markdown_text}
        - table: {type, headers, rows, caption}
        - figure: {type, image_path, caption, alt, figure_number}
        - code_block: {type, language, code}
        - quote: {type, text}
        - page_break: {type}
        - horizontal_rule: {type}
    """
    btype = block.get("type", "paragraph")

    if btype == "heading":
        _render_heading(doc, block)
    elif btype == "paragraph":
        _render_paragraph(doc, block)
    elif btype == "table":
        _render_table(doc, block)
    elif btype == "figure":
        _render_figure(doc, block)
    elif btype == "code_block":
        _render_code_block(doc, block)
    elif btype == "quote":
        _render_quote(doc, block)
    elif btype == "page_break":
        doc.add_page_break()
    elif btype == "horizontal_rule":
        _render_horizontal_rule(doc)
    else:
        # exhaust-retry: treat as paragraph
        _render_paragraph(doc, {"text": str(block.get("text", ""))})


def _render_heading(doc, block):
    """渲染标题，level 1-4 映射到 Word Heading 1-4"""
    level = block.get("level", 1)
    text = block.get("text", "")
    heading_level = min(max(level, 1), 4)

    p = doc.add_paragraph()
    if heading_level == 1:
        p.style = doc.styles['Heading 1']
    elif heading_level == 2:
        p.style = doc.styles['Heading 2']
    elif heading_level == 3:
        p.style = doc.styles['Heading 3']
    else:
        p.style = doc.styles['Heading 4']

    run = p.add_run(text)
    _set_run_font(run, size=Pt({1: 22, 2: 16, 3: 13, 4: 11}[heading_level]),
                  bold=True)


def _render_paragraph(doc, block):
    """渲染段落，支持内联格式（粗体、斜体、行内代码、链接）"""
    text = block.get("text", "")
    md_text = block.get("markdown_text", text)

    p = doc.add_paragraph()
    _set_paragraph_spacing(p, line_spacing=LINE_SPACING_BODY,
                           space_before=Pt(2), space_after=Pt(2))

    if not md_text:
        run = p.add_run(text)
        _set_run_font(run)
        return

    _parse_inline_markdown(p, md_text)


def _parse_inline_markdown(paragraph, text):
    """
    解析内联 Markdown 并添加到段落中。
    支持: **bold**, *italic*, `code`, [text](url), ~~strikethrough~~
    """
    # 正则：匹配内联元素
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)|'          # bold
        r'(\*(.+?)\*)|'              # italic
        r'(~~(.+?)~~)|'              # strikethrough
        r'(`(.+?)`)|'                # code
        r'(\[([^\]]+)\]\(([^)]+)\))'  # link
    )

    last_end = 0
    for m in pattern.finditer(text):
        start, end = m.start(), m.end()

        # 前面的纯文本
        if start > last_end:
            run = paragraph.add_run(text[last_end:start])
            _set_run_font(run)

        if m.group(1):   # bold
            run = paragraph.add_run(m.group(2))
            _set_run_font(run, bold=True)
        elif m.group(3):  # italic
            run = paragraph.add_run(m.group(4))
            _set_run_font(run, italic=True)
        elif m.group(5):  # strikethrough
            run = paragraph.add_run(m.group(6))
            _set_run_font(run)
            run.font.strike = True
        elif m.group(7):  # inline code
            run = paragraph.add_run(m.group(8))
            _set_run_font(run, name=FONT_MONO, size=Pt(9.5),
                          color=RGBColor(0xC7, 0x25, 0x4E))
        elif m.group(9):  # link
            link_text = m.group(10)
            link_url = m.group(11)
            run = paragraph.add_run(link_text)
            _set_run_font(run, color=RGBColor(0x05, 0x63, 0xC1))
            run.underline = True

        last_end = end

    # 剩余文本
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        _set_run_font(run)


def _render_table(doc, block):
    """
    渲染三线表。
    顶线 1.5pt，表头底线 0.75pt，底线 1.5pt，无竖线。
    """
    headers = block.get("headers", [])
    rows = block.get("rows", [])
    caption = block.get("caption", "")

    if not rows:
        return

    ncols = len(headers) if headers else max(len(r) for r in rows)
    nrows_data = len(rows)
    has_header = bool(headers)
    nrows_total = nrows_data + (1 if has_header else 0)

    # 标题（表N）
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cap.add_run(caption)
        _set_run_font(run, size=Pt(9), color=COLOR_CAPTION)
        _set_paragraph_spacing(p_cap, line_spacing=1.2,
                               space_before=Pt(6), space_after=Pt(2))

    table = doc.add_table(rows=nrows_total, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 填入表头
    if has_header:
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(h))
            _set_run_font(run, size=Pt(10), bold=True)

    # 填入数据行
    for i, row in enumerate(rows):
        target_row = table.rows[i + (1 if has_header else 0)]
        for j, val in enumerate(row):
            if j >= ncols:
                break
            cell = target_row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _set_run_font(run, size=Pt(10))

    # 应用三线表边框
    _apply_three_line_borders(table, has_header, nrows_total, ncols)

    # 表格后留空
    p_after = doc.add_paragraph()
    _set_paragraph_spacing(p_after, line_spacing=1.0, space_before=0, space_after=Pt(4))


def _apply_three_line_borders(table, has_header, nrows, ncols):
    """对表格应用三线表样式 —— 顶线、表头底线、底线"""
    border_top = {"sz": TABLE_BORDER_TOP_BOTTOM, "color": "000000"}
    border_bottom = {"sz": TABLE_BORDER_TOP_BOTTOM, "color": "000000"}
    border_header = {"sz": TABLE_BORDER_HEADER, "color": "000000"}

    for i in range(nrows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            edges = {}
            if i == 0:
                edges["top"] = border_top
            if i == 0 and has_header:
                edges["bottom"] = border_header
            if i == nrows - 1:
                edges["bottom"] = border_bottom
            if edges:
                _set_cell_border(cell, **edges)


def _render_figure(doc, block):
    """
    渲染图片。
    block: {image_path, caption, alt, figure_number}
    """
    image_path = block.get("image_path", "")
    caption = block.get("caption", "")
    figure_number = block.get("figure_number", "")
    alt = block.get("alt", "")

    # 图片
    if image_path and os.path.exists(image_path):
        try:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            # 限制最大宽度 14cm，保持比例
            run.add_picture(image_path, width=Cm(14))
            _set_paragraph_spacing(p_img, line_spacing=1.0,
                                   space_before=Pt(6), space_after=Pt(2))
        except Exception as e:
            p_err = doc.add_paragraph()
            run = p_err.add_run(f"[图片加载失败: {image_path}]")
            _set_run_font(run, color=RGBColor(0xCC, 0x00, 0x00), size=Pt(9))
    else:
        # 占位框
        p_placeholder = doc.add_paragraph()
        p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_placeholder.add_run(f"[图片: {alt or image_path}]")
        _set_run_font(run, color=RGBColor(0x99, 0x99, 0x99), size=Pt(9))
        _set_paragraph_spacing(p_placeholder, line_spacing=1.0,
                               space_before=Pt(6), space_after=Pt(2))

    # 图注
    cap_text = caption or alt
    if figure_number:
        cap_text = f"图{figure_number}: {cap_text}" if cap_text else f"图{figure_number}"

    if cap_text:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_cap.add_run(cap_text)
        _set_run_font(run, size=Pt(9), color=COLOR_CAPTION)
        _set_paragraph_spacing(p_cap, line_spacing=1.2,
                               space_before=Pt(2), space_after=Pt(8))


def _render_code_block(doc, block):
    """
    渲染代码块 —— 等宽字体、灰色背景。
    """
    code = block.get("code", "")
    language = block.get("language", "")

    # 语言标签
    if language:
        p_lang = doc.add_paragraph()
        run = p_lang.add_run(f"  {language}")
        _set_run_font(run, name=FONT_MONO, size=Pt(8),
                      color=RGBColor(0x99, 0x99, 0x99))
        _set_paragraph_spacing(p_lang, line_spacing=1.0,
                               space_before=Pt(4), space_after=0)
        _add_shading(p_lang, "EEEEEE")

    # 代码行
    lines = code.split("\n")
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        _set_run_font(run, name=FONT_MONO, size=FONT_SIZE_CODE,
                      color=COLOR_CODE_TEXT)
        _set_paragraph_spacing(p, line_spacing=LINE_SPACING_CODE,
                               space_before=0, space_after=0)
        _add_shading(p, "F2F2F2")

    # 底部留白
    p_end = doc.add_paragraph()
    _set_paragraph_spacing(p_end, line_spacing=1.0, space_before=0, space_after=Pt(4))


def _render_quote(doc, block):
    """
    渲染引用块 —— 灰色背景、左侧边框。
    """
    text = block.get("text", "")
    lines = text.strip().split("\n")

    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 左侧缩进
        pf = p.paragraph_format
        pf.left_indent = Cm(1.0)
        run = p.add_run(line)
        _set_run_font(run, size=Pt(10), italic=True,
                      color=RGBColor(0x55, 0x55, 0x55))
        _set_paragraph_spacing(p, line_spacing=1.3,
                               space_before=0, space_after=0 if i < len(lines) - 1 else Pt(4))
        _add_shading(p, "F5F5F5")
        _add_left_border(p, "999999", "12", 8)


def _render_horizontal_rule(doc):
    """渲染水平分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    _set_run_font(run, size=Pt(8), color=RGBColor(0xCC, 0xCC, 0xCC))
    _set_paragraph_spacing(p, line_spacing=1.0, space_before=Pt(6), space_after=Pt(6))


# ============================================================================
# 页眉页脚
# ============================================================================

def add_header_footer(doc, title):
    """
    添加页眉（报告名称）和页脚（X / Y 格式页码）。
    """
    for section in doc.sections:
        # --- 页眉 ---
        header = section.header
        header.is_linked_to_previous = False
        p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_header.clear()
        run = p_header.add_run(title or "报告")
        _set_run_font(run, size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
        _set_paragraph_spacing(p_header, line_spacing=1.0)

        # --- 页脚：X / Y ---
        footer = section.footer
        footer.is_linked_to_previous = False
        p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_footer.clear()

        # PAGE / NUMPAGES 域
        run1 = p_footer.add_run()
        fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1._r.append(fld_begin)

        run2 = p_footer.add_run()
        instr = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
        run2._r.append(instr)

        run3 = p_footer.add_run()
        fld_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fld_sep)

        run4 = p_footer.add_run("1")
        _set_run_font(run4, size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

        run5 = p_footer.add_run()
        fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fld_end)

        run6 = p_footer.add_run(" / ")
        _set_run_font(run6, size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

        run7 = p_footer.add_run()
        fld_begin2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run7._r.append(fld_begin2)

        run8 = p_footer.add_run()
        instr2 = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>'
        )
        run8._r.append(instr2)

        run9 = p_footer.add_run()
        fld_sep2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run9._r.append(fld_sep2)

        run10 = p_footer.add_run("1")
        _set_run_font(run10, size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

        run11 = p_footer.add_run()
        fld_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run11._r.append(fld_end2)

        _set_paragraph_spacing(p_footer, line_spacing=1.0)


# ============================================================================
# Markdown 解析器
# ============================================================================

def parse_markdown_blocks(md_path):
    """
    将 Markdown 文件解析为 block 列表。
    每个 block 是一个 dict，包含 type 及相关字段。
    """
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    blocks = []
    i = 0
    figure_counter = 0
    table_counter = 0

    while i < len(lines):
        line = lines[i]

        # 空行跳过
        if line.strip() == "":
            i += 1
            continue

        # Front matter (---)
        if line.strip() == "---" and i == 0:
            j = i + 1
            while j < len(lines) and lines[j].strip() != "---":
                j += 1
            i = j + 1
            continue

        # 代码块
        if line.strip().startswith("```"):
            language = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({
                "type": "code_block",
                "language": language,
                "code": "\n".join(code_lines)
            })
            continue

        # 标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # 水平线
        if re.match(r'^[-*_]{3,}\s*$', line):
            blocks.append({"type": "horizontal_rule"})
            i += 1
            continue

        # 引用
        if line.strip().startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:])
                i += 1
            blocks.append({"type": "quote", "text": "\n".join(quote_lines)})
            continue

        # 表格
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                # 解析表头
                headers = [h.strip() for h in table_lines[0].split("|")[1:-1]]
                # 跳过分隔行
                rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split("|")[1:-1]]
                    rows.append(cells)

                table_counter += 1
                blocks.append({
                    "type": "table",
                    "headers": headers,
                    "rows": rows,
                    "caption": f"表{table_counter}"
                })
            else:
                # 不够两行，穷尽尝试为段落
                for tl in table_lines:
                    blocks.append({"type": "paragraph", "text": tl})
            continue

        # 图片
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img_match:
            figure_counter += 1
            blocks.append({
                "type": "figure",
                "alt": img_match.group(1),
                "image_path": img_match.group(2),
                "caption": img_match.group(1),
                "figure_number": str(figure_counter)
            })
            i += 1
            continue

        # 普通段落（可能跨行）
        para_lines = []
        while i < len(lines) and lines[i].strip() != "" \
                and not re.match(r'^(#{1,6})\s+', lines[i]) \
                and not lines[i].strip().startswith("```") \
                and not re.match(r'^[-*_]{3,}\s*$', lines[i]) \
                and not lines[i].strip().startswith("> ") \
                and not (lines[i].strip().startswith("|") and "|" in lines[i]) \
                and not re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', lines[i]):
            para_lines.append(lines[i])
            i += 1

        if para_lines:
            md_text = "\n".join(para_lines)
            plain_text = re.sub(r'[*_`~\[\]\(\)#]', '', md_text).strip()
            blocks.append({
                "type": "paragraph",
                "text": plain_text,
                "markdown_text": md_text
            })

    return blocks


# ============================================================================
# 元数据提取
# ============================================================================

def extract_meta(md_path):
    """
    从 Markdown 文件头部的 YAML front matter 或首行标题提取元数据。
    """
    meta = {
        "title": "",
        "subtitle": "",
        "author": "",
        "version": "v1.0",
        "run_date": datetime.date.today().strftime("%Y-%m-%d")
    }

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")

    # 尝试 YAML front matter
    if lines and lines[0].strip() == "---":
        yaml_lines = []
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                break
            yaml_lines.append(lines[i])

        for yl in yaml_lines:
            kv = yl.split(":", 1)
            if len(kv) == 2:
                key = kv[0].strip().lower()
                val = kv[1].strip().strip('"').strip("'")
                if key == "title":
                    meta["title"] = val
                elif key == "subtitle":
                    meta["subtitle"] = val
                elif key == "author":
                    meta["author"] = val
                elif key == "version":
                    meta["version"] = val
                elif key == "date":
                    meta["run_date"] = val

    # exhaust-retry: 用第一个 # 标题
    if not meta["title"]:
        for line in lines:
            h1 = re.match(r'^#\s+(.+)$', line)
            if h1:
                meta["title"] = h1.group(1).strip()
                break

    # exhaust-retry: 用文件名
    if not meta["title"]:
        meta["title"] = os.path.splitext(os.path.basename(md_path))[0]

    return meta


# ============================================================================
# 主流程
# ============================================================================

def build_docx(md_path, output_path):
    """
    主函数：从 Markdown 构建 .docx。
    """
    print(f"[build_docx] 读取: {md_path}")

    # 提取元数据
    meta = extract_meta(md_path)
    print(f"[build_docx] 标题: {meta['title']}")

    # 解析 blocks
    blocks = parse_markdown_blocks(md_path)
    print(f"[build_docx] 解析到 {len(blocks)} 个内容块")

    # 创建文档（使用 reference.docx 模板以继承样式）
    _template_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'reference.docx')
    try:
        doc = Document(_template_path)
    except Exception as e:
        print(f"⚠️ 警告: reference.docx 模板加载失败（{e}），使用默认文档。建议在 Microsoft Word 中按 reference.docx 规范创建真实的 .docx 模板文件。")
        doc = Document()

    # 基础样式设置
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    sPr = style.element.get_or_add_rPr()
    sFonts = sPr.find(qn('w:rFonts'))
    if sFonts is None:
        from docx.oxml import OxmlElement
        sFonts = OxmlElement('w:rFonts')
        sPr.insert(0, sFonts)
    sFonts.set(qn('w:eastAsia'), FONT_BODY)

    # 封面
    add_cover(doc, meta)

    # TOC
    add_toc_field(doc)

    # 渲染内容
    for idx, block in enumerate(blocks):
        try:
            render_block(doc, block)
        except Exception as e:
            p_err = doc.add_paragraph()
            run = p_err.add_run(f"[渲染块 #{idx} ({block.get('type','?')}) 失败: {e}]")
            _set_run_font(run, color=RGBColor(0xCC, 0x00, 0x00), size=Pt(9))
            print(f"[build_docx] 警告: 块 #{idx} 渲染失败 - {e}", file=sys.stderr)

    # 页眉页脚
    add_header_footer(doc, meta["title"])

    # 保存
    doc.save(output_path)
    print(f"[build_docx] 已输出: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="从 Markdown 报告构建 .docx 文件（python-docx 穷尽尝试路径）"
    )
    parser.add_argument("input", help="输入 Markdown 文件路径")
    parser.add_argument("-o", "--output", help="输出 .docx 文件路径")
    args = parser.parse_args()

    md_path = args.input
    if not os.path.exists(md_path):
        print(f"错误: 文件不存在 - {md_path}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output
    if not output_path:
        base = os.path.splitext(os.path.basename(md_path))[0]
        output_dir = os.path.dirname(md_path) or "."
        output_path = os.path.join(output_dir, f"{base}.docx")

    build_docx(md_path, output_path)


if __name__ == "__main__":
    main()